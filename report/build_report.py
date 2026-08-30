#!/usr/bin/env python3
"""
Build 'Calorimeter study.docx' — appends the full technical body to the existing
Roboteq template, keeping its title page, revision table, TOC field, headers and
footers untouched.

Run:  python build_report.py
"""
from pathlib import Path
import copy
import re
import shutil

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path("/Users/vasilisgkionis/Desktop/RoboteQ/Calorimeter study.docx")
BAK = SRC.with_name(SRC.stem + " \u2014 pre-report backup.docx")
FIGDIR = Path.home() / "roboteq" / "python"
PAGE_W = Inches(6.27)          # A4 minus 1" margins each side

MONO = "Consolas"
MATH = "Cambria Math"

# Always build from the pristine template so re-running is safe.
if not BAK.exists():
    shutil.copy2(SRC, BAK)
    print(f"backup written: {BAK}")
doc = Document(str(BAK))

# ---- insertion cursor -----------------------------------------------------
# python-docx only appends.  ANCHOR lets a block be placed *before* an existing
# element instead, which is how the new 4.2-4.4 material gets in ahead of the
# "Topology evaluation" heading that was already in the template.
ANCHOR = [None]


def place(el):
    if ANCHOR[0] is not None:
        ANCHOR[0].addprevious(el)
    return el


def _p(style="Normal"):
    p = doc.add_paragraph(style=style)
    place(p._element)
    return p


def _t(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    place(t._tbl)
    return t

# ───────────────────────────────────────────────────────────── run/text engine

TOKEN = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\*[^*\n]+?\*|_\{[^}]*\}|\^\{[^}]*\}"
    r"|_[A-Za-z0-9]+|\^[A-Za-z0-9]+)")


def _emit(par, text, *, base_font=None, size=None, italic=False, bold=False,
          color=None):
    """Write `text` into `par`, honouring **bold**, `code`, _{sub} and ^{sup}."""
    for tok in TOKEN.split(text):
        if not tok:
            continue
        # bold / italic spans may themselves contain subscripts, so recurse
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            _emit(par, tok[2:-2], base_font=base_font, size=size,
                  italic=italic, bold=True, color=color)
            continue
        if (tok.startswith("*") and tok.endswith("*") and len(tok) > 2
                and not tok.startswith("**")):
            _emit(par, tok[1:-1], base_font=base_font, size=size,
                  italic=True, bold=bold, color=color)
            continue
        r = par.add_run()
        if tok.startswith("`") and tok.endswith("`") and len(tok) > 1:
            r.text = tok[1:-1]
            r.font.name = MONO
            r.font.size = Pt((size or 11) - 1.5)
        elif tok.startswith("_"):
            r.text = tok[2:-1] if tok.startswith("_{") else tok[1:]
            r.font.subscript = True
        elif tok.startswith("^"):
            r.text = tok[2:-1] if tok.startswith("^{") else tok[1:]
            r.font.superscript = True
        else:
            r.text = tok
        if r.font.name is None and base_font:
            r.font.name = base_font
        if size:
            r.font.size = Pt(size)
        if italic:
            r.italic = True
        if bold:
            r.bold = True
        if color:
            r.font.color.rgb = color
    return par


# ─────────────────────────────────────────────────────────────────── builders

def H(level, text):
    p = _p(style=f"Heading {level}")
    _emit(p, text)
    return p


def P(text="", *, size=None, space_after=6, italic=False, style="Normal"):
    p = _p(style=style)
    _emit(p, text, size=size, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def BUL(items, *, size=None, indent=0.25):
    for it in items:
        p = _p(style="Normal")
        _emit(p, "•\t" + it, size=size)
        pf = p.paragraph_format
        pf.left_indent = Inches(indent + 0.25)
        pf.first_line_indent = Inches(-0.25)
        pf.space_after = Pt(3)
        pf.space_before = Pt(0)


def NUM(items, *, size=None, indent=0.25):
    for i, it in enumerate(items, 1):
        p = _p(style="Normal")
        _emit(p, f"{i}.\t" + it, size=size)
        pf = p.paragraph_format
        pf.left_indent = Inches(indent + 0.3)
        pf.first_line_indent = Inches(-0.3)
        pf.space_after = Pt(3)
        pf.space_before = Pt(0)


def EQ(text, *, tag=None, size=12):
    p = _p(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _emit(p, text, base_font=MATH, size=size, italic=False)
    if tag:
        r = p.add_run("\t\t(" + tag + ")")
        r.font.size = Pt(size - 2)
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)
    return p


def MONOBLOCK(lines, *, size=8.5, indent=0.2):
    last = None
    for ln in lines:
        p = _p(style="Normal")
        last = p
        r = p.add_run(ln if ln else " ")
        r.font.name = MONO
        r.font.size = Pt(size)
        pf = p.paragraph_format
        pf.left_indent = Inches(indent)
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.0
    if last is not None:
        last.paragraph_format.space_after = Pt(8)


def _set_cell_width(cell, width):
    cell.width = width
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def TBL(headers, rows, weights=None, *, size=9, caption=None, cap_size=9):
    n = len(headers)
    t = _t(1, n)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    weights = weights or [1] * n
    total = sum(weights)
    widths = [Inches(PAGE_W.inches * w / total) for w in weights]

    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        _emit(c.paragraphs[0], h, size=size, bold=True)
        _shade(c, "DEEAF6")
        _set_cell_width(c, widths[j])
    _repeat_header(t.rows[0])

    for r in rows:
        cells = t.add_row().cells
        for j, v in enumerate(r):
            cells[j].text = ""
            _emit(cells[j].paragraphs[0], str(v), size=size)
            _set_cell_width(cells[j], widths[j])

    if caption:
        p = _p(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _emit(p, caption, size=cap_size, italic=True)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(10)
    else:
        _p().paragraph_format.space_after = Pt(4)
    return t


FIGN = [0]


def FIG(filename, caption, *, width=6.2):
    p = _p(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGDIR / filename), width=Inches(width))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    FIGN[0] += 1
    c = _p(style="Normal")
    c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _emit(c, f"**Figure {FIGN[0]}.** " + caption, size=9)
    c.paragraph_format.space_after = Pt(12)
    return FIGN[0]


def NOTE(title, text):
    """A shaded single-cell callout."""
    t = _t(1, 1)
    t.style = "Table Grid"
    t.autofit = False
    c = t.rows[0].cells[0]
    _set_cell_width(c, PAGE_W)
    c.text = ""
    _emit(c.paragraphs[0], f"**{title}**", size=9.5)
    p2 = c.add_paragraph()
    _emit(p2, text, size=9.5)
    p2.paragraph_format.space_after = Pt(2)
    _shade(c, "F2F2F2")
    _p().paragraph_format.space_after = Pt(4)


def PAGEBREAK():
    _p().add_run().add_break(6)  # WD_BREAK.PAGE == 6


# ───────────────────────────────────── 1. patch the existing tables in place

def table_after(heading_text, occurrence=0):
    """Return the nth table that follows the paragraph with this heading text."""
    body = list(doc.element.body.iterchildren())
    idx = None
    for i, el in enumerate(body):
        if el.tag == qn("w:p"):
            txt = "".join(n.text or "" for n in el.iter(qn("w:t")))
            if txt.strip() == heading_text.strip():
                idx = i
                break
    if idx is None:
        raise KeyError(heading_text)
    seen = 0
    for el in body[idx + 1:]:
        if el.tag == qn("w:tbl"):
            if seen == occurrence:
                for t in doc.tables:
                    if t._tbl is el:
                        return t
            seen += 1
    raise KeyError(f"no table {occurrence} after {heading_text!r}")


def fill_table(tbl, rows, *, size=9, weights=None):
    """Rewrite a 2-col template table: keep the header row, replace the rest."""
    for r in list(tbl.rows[1:]):
        r._tr.getparent().remove(r._tr)
    ncols = len(tbl.columns)
    weights = weights or [1] * ncols
    total = sum(weights)
    widths = [Inches(PAGE_W.inches * w / total) for w in weights]
    for j in range(ncols):
        _set_cell_width(tbl.rows[0].cells[j], widths[j])
    _repeat_header(tbl.rows[0])
    for r in rows:
        cells = tbl.add_row().cells
        for j, v in enumerate(r):
            cells[j].text = ""
            _emit(cells[j].paragraphs[0], str(v), size=size)
            _set_cell_width(cells[j], widths[j])


# ── Applicable documents ────────────────────────────────────────────────────
fill_table(table_after("Applicable documents"), [
    ["AD-01", "IEC 60034-2-1 — Rotating electrical machines: standard methods for "
              "determining losses and efficiency from tests"],
    ["AD-02", "IEC 61800-9-2 — Adjustable speed electrical power drive systems: "
              "ecodesign, energy efficiency indicators for power drive systems and "
              "motor starters"],
    ["AD-03", "Nidec Automation / Roboteq — Ultra-Low-Voltage Drive Catalog, "
              "17 Sep 2025 (F-series DUT class: 140 × 140 × 25 mm, 2 × 60 A / 1 × 120 A)"],
], weights=[1, 6])

# ── Reference documents ─────────────────────────────────────────────────────
fill_table(table_after("Reference Documents"), [
    ["RD-01", "Christen, D., Badstübner, U., Biela, J., Kolar, J. W., "
              "“Calorimetric Power Loss Measurement for Highly Efficient Converters”, "
              "IPEC 2010, ETH Zürich Power Electronic Systems Laboratory. "
              "Built air-gap double-jacket calorimeter, ±0.4 W / 1 %. "
              "Primary architecture reference."],
    ["RD-02", "Simpson, N., Hopkins, A. N., “An Accurate and Flexible Calorimeter "
              "Topology for Power Electronic System Loss Measurement”, IEMDC 2017, "
              "University of Bristol. Surface-heated closed calorimeter, < 1 % over "
              "20–100 W. Guard-actuator and commissioning reference."],
    ["RD-03", "Nair, D. G., “Calorimetric Power Loss Measurement of a Small Power "
              "Converter”, MSc thesis, Aalto University, 2013. Source of the "
              "double-jacket schematic and of the guard criterion "
              "|T_i − T_e| < P_{acc}·R_{th(in)} (§3.4.2). Also the unguarded "
              "counter-example."],
    ["RD-04", "Kosonen, A., et al., “Calorimetric concept for measurement of power "
              "losses up to 2 kW in electric drives”, IET Electric Power Applications, "
              "2013. Modularity across DUT ratings."],
    ["RD-05", "Malliband, P. D., Carter, B., Gordon, M., Warne, D. F., “Accurate "
              "measurement of induction motor losses using balance calorimeter”, "
              "Proc. IEE, vol. 138, no. 5, 1991."],
    ["RD-06", "Malliband, P. D., et al., “Design of a double-jacketed, closed type "
              "calorimeter for direct measurement of motor losses”, 7th Int. Conf. on "
              "Power Electronics and Variable Speed Drives, no. 456, 1998. "
              "Original DJC proposal."],
    ["RD-07", "Blaabjerg, F., Pedersen, J. K., Ritchie, E., “Calorimetric measuring "
              "systems for characterizing high frequency power losses in power "
              "electronic components and systems”, IEEE IAS Conf., vol. 2, "
              "pp. 1369–1376, 2002. First proposal of active air-gap temperature "
              "control."],
    ["RD-08", "Xiao, C., Chen, G., Odendaal, W. G. H., “Overview of Power Loss "
              "Measurement Techniques in Power Electronic Systems”, IEEE Trans. "
              "Industry Applications, pp. 657–664, 2007."],
    ["RD-09", "Jalilian, A., Gosbell, V. J., Perera, B. S. P., Cooper, P., “Double "
              "chamber calorimeter (DCC): a new approach to measure induction motor "
              "harmonic losses”, IEEE Trans. Energy Conversion, vol. 14, no. 3, 1999. "
              "The series calorimeter — distinct from the DJC."],
    ["RD-10", "“Overview of calorimetric systems used in loss determination of "
              "electric motors and drives”, IEEE, 2017. Source of the four-way "
              "open / closed / balanced / series classification."],
    ["RD-11", "Internal working notes, Obsidian vault "
              "`Roboteq Internship/Calorimetry/` — Study, thermo modeling, "
              "Parameter Calculation, Biot check, System Modeling Final, BOM manual, "
              "Zephyr map."],
    ["RD-12", "Simulation source, `~/roboteq/python/` — `calorimetry-ss.py` (plant), "
              "`calorimetry_sim.py` (closed loop, probes, plots), `params.yaml`."],
], weights=[1, 6])

# ── Abbreviations ───────────────────────────────────────────────────────────
abbr = [
    ["ADC", "Analog-to-Digital Converter"],
    ["AW", "Anti-Windup"],
    ["B1", "Measurement boundary — the surface enclosing the inner chamber contents"],
    ["BOM", "Bill of Materials"],
    ["DAQ", "Data Acquisition"],
    ["DC", "Direct Current — here also the zero-frequency gain of a transfer function"],
    ["DJC", "Double-Jacket Chamber"],
    ["DUT", "Device Under Test"],
    ["EPP", "Expanded Polypropylene"],
    ["ESD", "Electrostatic Discharge"],
    ["FF", "Feedforward"],
    ["HX", "Heat Exchanger"],
    ["I²C", "Inter-Integrated Circuit (two-wire serial bus)"],
    ["IMC", "Internal Model Control — the λ-tuning rule used for the guard PI"],
    ["LTI", "Linear Time-Invariant"],
    ["MCU", "Microcontroller Unit"],
    ["NA", "Not Applicable"],
    ["PI / PID", "Proportional-Integral / Proportional-Integral-Derivative controller"],
    ["PWM", "Pulse Width Modulation"],
    ["RK4", "Fourth-order Runge-Kutta fixed-step integrator"],
    ["RSS", "Root Sum of Squares"],
    ["RTD", "Resistance Temperature Detector (Pt100 / Pt1000)"],
    ["SISO", "Single-Input Single-Output"],
    ["TIM", "Thermal Interface Material"],
    ["ULV", "Ultra-Low Voltage (≤ 60 V — the Roboteq / Nidec Automation drive class)"],
    ["XPS", "Extruded Polystyrene (rigid foam insulation)"],
    ["ZOH", "Zero-Order Hold"],
]
abbr_tbl = table_after("Abbreviations And Acronyms ")
for r in list(abbr_tbl.rows):
    r._tr.getparent().remove(r._tr)
w = [Inches(PAGE_W.inches * 1 / 5), Inches(PAGE_W.inches * 4 / 5)]
for a, b in abbr:
    cells = abbr_tbl.add_row().cells
    for j, v in enumerate((a, b)):
        cells[j].text = ""
        _emit(cells[j].paragraphs[0], v, size=9, bold=(j == 0))
        _set_cell_width(cells[j], w[j])

# ── revision history ───────────────────────────────────────────────────────
rev = table_after("Revision History")
for r in rev.rows[1:]:
    if not "".join(c.text for c in r.cells).strip():          # first blank row
        vals = ["1.1", "30\u201308\u20132026", "G. Vasileios",
                "Sections 2\u20137 completed: reference documents, calorimetry "
                "theory, topology trade study and selection, design decisions, "
                "thermal model, state-space model and guard control design, "
                "Python simulation and results. Sections 8\u20139 (hardware, "
                "firmware) still to be written."]
        for c, v in zip(r.cells, vals):
            c.text = ""
            _emit(c.paragraphs[0], v, size=9)
        break

# ── ask Word to refresh the table of contents on open ──────────────────────
_settings = doc.settings.element
if _settings.find(qn("w:updateFields")) is None:
    _uf = OxmlElement("w:updateFields")
    _uf.set(qn("w:val"), "true")
    _settings.append(_uf)

# ── drop the "Subtopic / Subtopic / Content" placeholders ───────────────────
for p in list(doc.paragraphs):
    if p.text.strip() in ("Subtopic", "Content") and p.style.name in (
            "Heading 3", "Heading 4", "Normal"):
        p._element.getparent().remove(p._element)

print("template patched; appending body…")

# ═══════════════════════════════════════════════════════════════════════════
#                            THE BODY
# ═══════════════════════════════════════════════════════════════════════════

# The template already contains §4.1 "Calorimeter" and the opening paragraph of
# §4.5 "Topology evaluation".  Sections 4.2–4.4 have to be inserted *before* that
# heading; everything from 4.5.1 onwards is appended after its paragraph.

def _heading_element(text):
    for el in doc.element.body.iterchildren():
        if el.tag == qn("w:p"):
            t = "".join(n.text or "" for n in el.iter(qn("w:t")))
            if t.strip() == text:
                return el
    raise KeyError(text)


ANCHOR[0] = _heading_element("Topology evaluation")

# ---------------------------------------------------------------- 4.2 ------
H(2, "Why calorimetry, and not an electrical measurement")

P("The direct way to find a converter's loss is to measure input and output "
  "power electrically and subtract. That method degrades exactly where this "
  "project needs it most, because the measurand is the small difference of two "
  "large, nearly equal numbers.")

P("For the F-series DUT at 1.5 kW, the loss to be resolved is between 15 W "
  "(η = 99 %) and 200 W (η ≈ 88 %). Resolving the 15 W point to ±1 W by "
  "subtraction requires each of the two 1.5 kW measurements to be accurate to "
  "better than ±0.7 W, i.e. ±0.05 % of reading. Real bench instrumentation does "
  "not do this:")

TBL(
    ["Power-analyser class", "Uncertainty per channel at 1.5 kW",
     "Uncertainty on P_{in} − P_{out} (RSS)", "As a fraction of a 15 W loss"],
    [["±0.05 % of reading", "0.75 W", "1.06 W", "7.1 %"],
     ["±0.1 %", "1.5 W", "2.12 W", "14.1 %"],
     ["±0.2 %", "3.0 W", "4.24 W", "28.3 %"],
     ["±0.5 %", "7.5 W", "10.6 W", "70.7 %"]],
    weights=[2, 2, 2.2, 1.8],
    caption="Table 1. Why subtraction fails at the high-efficiency end. The "
            "uncertainty is fixed by the throughput, not by the loss, so it does "
            "not shrink as the DUT gets better. The same instrument that reads a "
            "200 W loss to 2 % reads a 15 W loss to 14 % or worse.")

P("Calorimetry does not have this structure. It measures the dissipated heat "
  "itself, so its uncertainty is set by the flow meter, the temperature-pair "
  "match and the guard residual — all of which are roughly **constant in watts** "
  "across the range, and none of which depend on the DUT's throughput, waveform, "
  "switching frequency, power factor or EMI. A calorimeter that is good to ±1 W "
  "is good to ±1 W at 15 W and at 200 W alike.")

TBL(["Aspect", "Electrical (input − output)", "Calorimetric"],
    [["What is measured", "Two large electrical powers, then their difference",
      "The heat leaving the DUT, directly"],
     ["Dominant error", "Difference of two similar large numbers — error grows "
      "as a fraction of the loss as efficiency rises",
      "Absolute — stays roughly constant in watts"],
     ["Sensitive to", "Waveform distortion, phase/timing error, sensor bandwidth, EMI",
      "Ambient stability, wall leakage, calibration drift"],
     ["Electrical connection to DUT", "Required", "None — fully decoupled"],
     ["Speed", "Seconds", "Minutes to hours per point (thermal settling)"],
     ["Best suited to", "Low to moderate efficiency, transient measurement",
      "Verifying loss on high-efficiency devices (η > ~95 %)"]],
    weights=[1.5, 3, 3],
    caption="Table 2. The two methods compared. Calorimetry is slower and needs "
            "more hardware; it buys accuracy that does not degrade with the DUT's "
            "efficiency.")

P("Published, well-built calorimeters report about **±0.4 W absolute, or 0.2–1 % "
  "of full scale** (RD-01, RD-02), essentially flat across the power range. The "
  "accuracy target adopted for this rig is **P_{acc} = ±1 W**, chosen as a "
  "realistic first-build figure with margin against the literature best. At "
  "1.5 kW throughput, ±1 W resolves 0.07 efficiency points — enough to separate "
  "97.0 % from 97.5 % (a 7.5 W step) without ambiguity.")

# ---------------------------------------------------------------- 4.3 ------
H(2, "Calorimeter families")

P("Calorimeters built for power-loss measurement are classified into four "
  "chamber types (RD-10). Each can be run air- or liquid-cooled.")

TBL(["Type", "How it works", "Trade-off"],
    [["Open", "Single pass of coolant drawn from ambient through the chamber; "
      "loss from inlet/outlet ΔT and flow",
      "Simplest build; accuracy is at the mercy of ambient temperature and flow "
      "fluctuations"],
     ["Closed", "DUT sealed in an insulated enclosure with a closed coolant loop "
      "that dumps to an external exchanger",
      "Much better ambient isolation; usually built with a guarded (double-jacket) "
      "wall to stop the residual leak"],
     ["Balanced (substitution)",
      "A trim heater holds the chamber at a fixed balance point; loss = the "
      "**reduction** in trim-heater power relative to a DUT-off baseline",
      "Very accurate and needs no flow measurement, but requires a control loop "
      "and a baseline calibration run"],
     ["Series", "Two chambers share one coolant loop — one holds a reference "
      "heater of known power, the other the DUT — so the rig self-calibrates "
      "continuously",
      "Best drift immunity, at the cost of a second full chamber"]],
    weights=[1.4, 3.3, 2.8],
    caption="Table 3. The four chamber families. This project builds a **closed, "
            "guarded, liquid-cooled** chamber with a flow readout, and keeps "
            "substitution as the calibration method rather than as the readout.")

NOTE("Two readout philosophies, kept separate",
     "A guarded chamber can report the loss in two different ways, and the "
     "project notes use both words. **Flow readout** (this build): the loss is "
     "ṁ·c_p·ΔT_w through the water loop, and the guard exists only to null the "
     "wall so that reading is complete. **Substitution readout**: the loss is a "
     "difference of trim-heater powers, and no flow measurement is needed. This "
     "rig uses the flow readout for the measurement and the substitution "
     "principle only for commissioning — the inner-chamber resistor injects a "
     "known electrical power that the flow readout must reproduce.")

# ---------------------------------------------------------------- 4.4 ------
H(2, "Reference builds")

P("Four published builds set the design envelope and supplied the specific "
  "techniques that were copied. Three are used as positive references; the "
  "fourth is the counter-example that justifies the guard.")

TBL(["Source", "Topology", "DUT / loss range", "Reported accuracy",
     "What was taken from it"],
    [["**ETH Zürich, IPEC 2010** (RD-01)",
      "Air-gap double-jacket, actively servoed gap",
      "5 kW DC-DC at 99.1 % η; also a 1.6 kW PFC at 15 W loss",
      "±0.4 W or 1 %, roughly flat over 10–100 W",
      "The architecture itself: nested shells, servoed air gap, closed water loop "
      "with its own radiator and fans, in-chamber preheat to shorten settling, "
      "digital I²C wall sensors"],
     ["**Bristol, IEMDC 2017** (RD-02)",
      "Closed single box with a PI-controlled copper-tape surface heater — the "
      "guard collapsed onto the wall",
      "20–100 W (predicted ceiling ≈ 250 W)",
      "< 1 % at all 12 load points; < 0.4 W absolute",
      "The guard as one software PI loop; in-situ gravimetric flow calibration of "
      "a cheap flow sensor; automatic steady-state detection from the temperature "
      "slope; commodity PC-watercooling parts; single-point calibration"],
     ["**Kosonen et al., IET EPA 2013** (RD-04)",
      "One rig covering several drive ratings",
      "Up to 2 kW of loss",
      "—",
      "The modularity argument: build one chamber and re-range it per DUT class "
      "rather than one rig per product"],
     ["**Nair, Aalto MSc 2013** (RD-03)",
      "Closed **single-wall, unguarded** chamber (a DJC was reviewed, then not "
      "built)",
      "0.75 kW converter at ~97 % η",
      "±1.5 % below 50 W, **degrading with power**",
      "The counter-example. Its own conclusion names wall leakage as the dominant "
      "error. Also the source of the DJC schematic and of the guard criterion "
      "used in §4.6.3"]],
    weights=[1.5, 2, 1.7, 1.6, 3.2], size=8.5,
    caption="Table 4. The four reference builds. Nair and ETH differ almost only "
            "in whether the wall gradient is actively nulled, and their accuracy "
            "figures differ by roughly a factor of four — which is the whole "
            "argument for the guard.")

P("Two of these builds, a decade apart and unrelated, independently converged on "
  "the same commodity parts: XPS/EPP foam shells, Aqua Computer PC-watercooling "
  "radiators, a peristaltic pump, a cheap flow sensor rescued by gravimetric "
  "calibration, and Class-A platinum RTDs. That convergence is why this project's "
  "BOM starts from the same categories rather than from laboratory-instrument "
  "catalogues.")

# ---------------------------------------------------------------- 4.5 ------
# Back to appending: everything below follows the template's own
# "Topology evaluation" heading and its opening paragraph.
ANCHOR[0] = None

P("Those three are compared in §4.5.1; the reasoning that led to them, starting "
  "from a wider four-way trade study, is in §4.5.2, and the decision in §4.5.3. "
  "The remainder of this section describes the selected topology and the design "
  "choices made during the study. Sections 5 to 7 then develop the model of that "
  "topology in increasing detail: the thermal network (§5), its state-space form "
  "and the control design that follows from it (§6), and the Python simulation "
  "used to size the actuators and predict the settling schedule (§7).")

H(3, "The three candidate topologies")

TBL(["", "**1 — Single wall, passive**", "**2 — Surface-heated single box**",
     "**3 — Full air-gap double jacket**"],
    [["Shells to build", "1", "1", "2, nested"],
     ["Guard actuator", "None — insulation only",
      "Copper-foil strip heater bonded to the outer surface, PI-controlled",
      "Heated and fanned air gap between the shells, PI-controlled"],
     ["Guard sensing", "n/a", "Surface RTDs, ~1 mm under the skin",
      "Gap-air temperature T_e vs chamber air T_i"],
     ["Feedthroughs cross", "1 wall", "1 wall", "2 walls"],
     ["Wall leak", "Slowed, never cancelled; drifts with ambient",
      "Actively nulled at the wall surface",
      "Actively nulled across the whole inner wall"],
     ["Demonstrated accuracy",
      "±1.5 %, worsening with power (RD-03)",
      "< 1 % over 20–100 W (RD-02)", "±0.4 W flat (RD-01)"],
     ["Estimated cost", "Lowest", "≈ €620–1 020", "≈ €710–1 150"],
     ["Main risk", "Accuracy — will not meet ±1 W",
      "Heater and temperature uniformity over a small box",
      "Build effort: a second shell to fabricate and seal"]],
    weights=[1.5, 2.3, 2.6, 2.4], size=8.5,
    caption="Table 5. The three topologies. The measurement core — water loop, "
            "flow meter, ΔT pair, in-chamber radiator and fans — is identical in "
            "all three and costs about €490–750; the whole difference is the guard.")

H(3, "How the shortlist was reached")

P("The study did not start with these three. The first pass, before the "
  "literature review, compared four configurations against cost, build time, "
  "thermal time constant and accuracy:")

TBL(["", "**A — Open, low-cost**", "**B — Closed, passive guard**",
     "**C — Closed, active guard**", "**D — Repurposed chamber**"],
    [["Readout", "Single-pass air, ṁ·c_p·ΔT",
      "Recirculating air in an insulated box, ṁ·c_p·ΔT",
      "Trim-heater power difference (substitution)",
      "Same as B, retrofitted into an existing enclosure"],
     ["Parts cost", "~$100–150", "~$500–950", "~$450–800",
      "~$530–810 if a chamber is on hand; $1 800–3 300 otherwise"],
     ["Accuracy at 150–200 W", "±10–20 %", "±3–5 %", "±1–2 %", "±3–7 %"],
     ["Accuracy at 15–30 W", "±10–20 %", "±15–30 %", "±1–2 %", "±10–20 %"],
     ["Build effort", "~1 week", "~2–3 weeks", "~4–6 weeks",
      "~1–2 weeks if a chamber exists"],
     ["Outcome", "Kept only as a week-1 shakedown of the sensor and logging chain",
      "Fallback if schedule ran out", "**Selected**",
      "**Ruled out on day 1** — no chamber exists at Roboteq HQ to repurpose"]],
    weights=[1.3, 1.9, 2.2, 2.3, 2.3], size=8.5,
    caption="Table 6. The original four-way trade study. Cost figures are 2026 "
            "order-of-magnitude estimates for individual components, not vendor "
            "quotes.")

P("Two findings decided it. First, **B's weakness is at the wrong end of the "
  "range**: a flow sensor's and an RTD pair's uncertainties are fixed in absolute "
  "terms, so at the 200 W / 10 K design point they are a small fraction of the "
  "reading, but at 15 W and 0.75 K of air-side rise the same fixed ±0.15 K is "
  "15–20 % of the signal. Relative accuracy therefore degrades exactly at the "
  "high-efficiency operating points calorimetry exists to serve. Second, once the "
  "DUT's real loss range was known, **C was no longer the more expensive option** "
  "— it removes the precision flow sensor that dominates B's cost and replaces it "
  "with foam, a resistor and a control loop. C dominated B on both cost and "
  "accuracy; the only axis where B still won was schedule risk.")

P("The literature review then reshaped what \"active guard\" means in hardware. "
  "The Bristol paper (RD-02) showed the same guard physics implemented as a "
  "heater bonded to the outside of a **single** box, which removes the second "
  "shell — the main source of C's schedule risk. That turned the original A/B/C/D "
  "question into the three-way comparison of Table 5, all three columns of which "
  "come from real, published, measured builds rather than from estimates.")

H(3, "Decision")

P("**Topology 3 — the full air-gap double-jacket chamber — is the build**, "
  "confirmed 2026-07-10 and reaffirmed 2026-07-13. Topology 2 stays documented as "
  "the fallback if the air-gap guard underperforms. The reasoning:")

BUL([
    "**A passive wall cannot meet ±1 W.** Worked for this rig's own geometry: at "
    "a chamber 10 K above ambient and R_{th(in)} ≈ 1.8 K/W, an unguarded inner "
    "wall leaks about 5.6 W. Even a calibrated leak model leaves a residual from "
    "ambient drift, draught-driven film-coefficient variation and seam ageing — "
    "realistically ~10 % of the correction, i.e. 0.5–0.6 W, before any other "
    "error term is counted. Nair (RD-03) measured exactly this failure mode.",
    "**The ETH air-gap variant is the best-characterised precedent** at our "
    "accuracy target and is the one whose schematic, criterion and preheat trick "
    "the design already follows.",
    "**Its cost premium over the surface-heated variant is small** (≈ €90–130) "
    "and its main risk — fabricating and sealing a second shell — is largely a "
    "purchasing question, since insulated foam boxes are a commodity product that "
    "can be ordered close to size rather than built from sheet stock.",
    "**Uniformity risk favours the air gap on a box this size.** The Bristol "
    "surface heater was validated on a 525 × 305 × 195 mm box with all faces "
    "driven in series, and per-face control is listed there as future work. On a "
    "smaller chamber, a stirred air gap distributes the guard's heat for us, "
    "which a bonded foil heater does not.",
])

# ---------------------------------------------------------------- 4.6 ------
H(2, "Selected topology — the double-jacket chamber")

H(3, "Architecture")

P("Two nested insulated shells. The inner chamber holds the DUT, the water-loop "
  "pickup radiator, circulation fans, a preheat/calibration resistor and the "
  "chamber temperature sensor. The gap between the shells has its own heater and "
  "fan and is servoed to track the inner chamber. The water loop crosses both "
  "walls and is the only intended exit for the DUT's heat.")

MONOBLOCK([
    "        Ta — room / ambient, free to drift",
    "  ┌────────────────────────────────────────────────────────────┐",
    "  │  Outer shell — insulated, heated and fanned                │",
    "  │  air gap servoed to track the inner chamber                │",
    "  │                                                            │",
    "  │        Te — gap air        Pheater(e), Pfan(e)             │",
    "  │   ┌────────────────────────────────────────────────┐       │",
    "  │   │  Inner chamber                    B1           │       │",
    "  │   │                                                │       │",
    "  │   │   water in  ──►[ pickup radiator ]──► water out│       │",
    "  │   │                        ▲                       │       │",
    "  │   │                     [ DUT ] ──► Ploss          │       │",
    "  │   │                                                │       │",
    "  │   │   Ti — chamber air     Pheater(i), Pfan(i)     │       │",
    "  │   └────────────────────────────────────────────────┘       │",
    "  │      inner wall:  Rth(in) = d /(λ·A) ,  Ggap               │",
    "  └────────────────────────────────────────────────────────────┘",
    "         outer wall:  Rth(out) ,  Gout        → to ambient",
], size=8.5)

P("The labelling follows RD-03 Fig. 3.4 exactly: **T_i** inner chamber, **T_e** "
  "gap air, **T_a** ambient, with a heater and a fan on each side of the inner "
  "wall. Three boundaries matter, and every equation in §5 is a first-law balance "
  "across one of them:")

TBL(["Boundary", "Encloses", "What crosses it"],
    [["**B1**", "The inner-chamber contents — the measurement boundary",
      "In: the DUT's loss (as electricity) and the metered auxiliaries. "
      "Out: the water's enthalpy rise (measured) and the wall leak (the error "
      "term the guard deletes)."],
     ["**B2**", "The air gap — the guard's boundary",
      "In: gap heater and gap fan power, plus whatever leaks in from B1 "
      "(ideally zero). Out: loss to ambient through the outer wall — unmetered, "
      "and harmless, because it is not on the measured path."],
     ["**B3**", "Any single lump inside, when its own temperature is needed",
      "Used for the foam wall and the DUT body only."]],
    weights=[1, 2.6, 5],
    caption="Table 7. The three boundaries.")

H(3, "Working principle")

NUM([
    "The DUT is mounted in the inner chamber and run at its operating point. All "
    "of its loss becomes heat inside B1.",
    "The guard loop drives the gap air to the chamber's own temperature, so the "
    "gradient across the inner wall goes to zero.",
    "With no gradient, Fourier's law gives no conduction through that wall — the "
    "leak is cancelled, not merely slowed. Radiation across the gap is nulled by "
    "the same condition.",
    "The only remaining exit for the heat is the water loop.",
    "The loss is read from the coolant's mass flow and its inlet-to-outlet "
    "temperature rise, with the metered in-chamber auxiliaries subtracted.",
])

EQ("P_{loss} = ṁ · c_{p,w} · ΔT_w − P_{aux(i)}", tag="1")

P("This is the whole instrument. Everything else in this report exists either to "
  "justify the two conditions in steps 2 and 3, or to price how imperfectly they "
  "are met.")

H(3, "The guard criterion and the two gates")

P("Equation (1) is only valid while two conditions hold. Both are checked in "
  "firmware before a point is logged, and both are stated as explicit gates.")

P("**Null gate.** From RD-03 §3.4.2, the residual gradient allowed across the "
  "inner wall is fixed by the accuracy budget and the wall's own resistance:")

EQ("| T_i − T_e |  <  P_{acc} · R_{th(in)}", tag="2")

P("With P_{acc} = 1 W and the design wall (XPS, 50 mm, λ ≈ 0.035–0.04 W/m·K over "
  "≈ 0.6 m² ⇒ G_{gap} ≈ 0.56 W/K, R_{th(in)} ≈ 1.79 K/W), the gate is "
  "**|T_i − T_e| < 1.79 K**. This is a design spec, not a comparison against the "
  "unguarded case: decide how accurate the rig must be, and the wall's resistance "
  "converts that directly into how tightly the servo has to hold the null. A "
  "realistic servo residual of ±0.3 K costs 0.17 W.")

P("**Steady gate.** The storage term must be dead, or the chamber's thermal mass "
  "is silently charging or discharging into the reading:")

EQ("| dP̂_{meas}/dt |  <  P_{thr} / τ_{dom}", tag="3")

P("The slope threshold is the residual power threshold divided by the slowest "
  "time constant the reading actually shows. It is computed per operating point "
  "rather than hard-coded, because τ_{dom} moves by a factor of about three "
  "across the power range. Equivalently, in temperature terms, holding a 0.2 W "
  "storage slice against C_i + C_w ≈ 2.45 kJ/K means |dT/dt| < 0.29 K/h.")

NOTE("Gate the reading, not the chamber temperature",
     "T_i is dominated by the fast air/water exchange mode and settles long "
     "before the reading does. A steady-state detector that watches T_i will go "
     "green while the water side is still transient. The gate must watch "
     "P̂_{meas} itself.")

P("A useful property of the architecture: **the outer shell's leak to the room "
  "does not need to be controlled at all.** The servo nulls one boundary only — "
  "chamber to gap. Whatever the gap then exchanges with the room never crosses "
  "the measured boundary, so those watts are accuracy-free. The room only has to "
  "be stable enough for the inner loop to track. This is the same principle as "
  "the guard ring in guarded-hot-plate conductivity testing (ASTM C177).")

# ---------------------------------------------------------------- 4.7 ------
H(2, "Design decisions taken during the study")

P("The design changed several times over the two months as specifications "
  "firmed up and as parts were confirmed. The table records what changed, when, "
  "and why, so that the working notes can be read in the right order.")

TBL(["Decision", "Date", "Outcome", "Reason"],
    [["Repurpose an existing chamber", "Day 1",
      "**Ruled out**", "No chamber, incubator or environmental enclosure exists "
      "at Roboteq HQ. The rig is a ground-up build."],
     ["Coolant medium", "Early July", "**Water, closed recirculating loop**",
      "Water's volumetric heat capacity is ~3 000–4 000× air's, so the same heat "
      "is carried by a manageable flow with a readable ΔT. Matches RD-01; RD-03's "
      "mains-fed open loop is explicitly what its own future-work section "
      "recommends moving away from."],
     ["Chamber topology", "2026-07-10, reaffirmed 07-13",
      "**Air-gap DJC (topology 3)**",
      "Accuracy at the low-loss end. The surface-heated variant (topology 2) "
      "stays documented as the fallback."],
     ["Readout method", "Early July", "**Flow readout, guard used only to null**",
      "Substitution keeps its role as the calibration method, not as the readout."],
     ["Preheat stage", "Mid July", "**Adopted**",
      "RD-01's technique: pre-load the chamber to near the expected operating "
      "point with a known resistor before the DUT is switched on, which removes "
      "most of the first-point warm-up. The same resistor doubles as the "
      "calibration reference."],
     ["Gap skin finish", "Mid July", "**Aluminium-foil lining, both facing skins**",
      "Radiation across the gap, not conduction, is the dominant leak with bare "
      "skins — see §5.4. Foil is a ~30× reduction and is what makes the guard "
      "spec affordable."],
     ["DUT-class modularity — copper cold block",
      "Designed July; **dropped 2026-08-27**",
      "**Removed from build scope**",
      "An optional water-cooled copper block was designed so a high-TDP DUT could "
      "dump heat straight into the water instead of into the chamber air. It was "
      "proven accuracy-neutral, but the F-series DUT dissipates about 15 W at its "
      "efficient operating points, where the block buys nothing. It was dropped "
      "as a time-box call, not a technical failure. See the note below."],
     ["Chamber temperature target and flow band", "2026-08-27",
      "**Chamber T_{ss} raised to 50–60 °C**",
      "Supervisor-driven. Raises the available ΔT and lowers the required flow. "
      "The consequences for the low-power corner are quantified in §7.9.2 and are "
      "not entirely benign."],
     ["Pump", "2026-08-27", "**JwardTech 304K/BT peristaltic, stepper variant**",
      "Replaces the previously listed Kamoer KPHM600, which at 600–800 mL/min was "
      "about 10× too high-flow once the target ΔT rose. The stepper's low end "
      "goes cleanly toward zero, unlike a brushless pump's duty-cycle deadband."],
     ["Flow metering", "2026-08-27", "**Biotech turbine sensor is the meter**",
      "Pump-as-flow-meter (step count × displacement) is not used for metrology; "
      "the pump is purely the flow actuator."],
     ["Chamber and wall sensing bus", "2026-08-27",
      "**TMP117 over I²C for both chamber and guard array**",
      "Resolves an open three-way choice. Confirms I²C as the sensor bus, and "
      "therefore the need for an address multiplexer, since the guard array alone "
      "needs 8–12 sensors and the part has only four addresses."],
     ["Water ΔT sensing", "2026-08-27", "**Jumo VIBROtemp pair, bare RTD**",
      "The ordered variant has no integrated transmitter, so a MAX31865-class "
      "front end per probe is required on the carrier board."],
     ["MCU and RTOS", "August", "**ESP32-S3 DevKitC running Zephyr**",
      "Native Wi-Fi for live telemetry, existing Zephyr board support, and a "
      "module that plugs into a carrier board rather than a bespoke SoC design."],
     ["Supply rail", "August", "**Single 12 V rail (LRS-350-12 class)**",
      "Sized bottom-up from the two 100 W heaters, fans and pump — about 22 A "
      "against the supply's 29 A."]],
    weights=[1.9, 1.2, 1.9, 5], size=8.5,
    caption="Table 8. Design decisions and revisions, in the order they were made.")

NOTE("On the copper block, and on notes that still mention it",
     "The rig was originally intended to carry an optional copper cold plate "
     "bolted to the DUT's baseplate, so that a high-TDP device could send most of "
     "its heat straight into the water (a split fraction α ≈ 0.97) rather than "
     "into the chamber air (α = 0). The analysis showed the reading is exactly "
     "**independent of α** — see §6.6 — so the block was never an accuracy "
     "feature; it was a way to widen the usable DUT range past ~200 W. Since the "
     "F-series DUT dissipates roughly 15 W at its efficient operating points and "
     "200 W at worst, the block earns nothing on this DUT, and it was dropped "
     "from build scope on 2026-08-27 with the supervisor's agreement. **The "
     "built rig is the α = 0, radiator-only case.** Several working notes, some "
     "BOM links and the default `alpha: 0.97` in `params.yaml` still refer to the "
     "block; they are historical and are not part of the delivered design.")

# ═════════════════════════════════════════════════════════ 5. thermal model
H(1, "Thermal Model")

H(2, "Modelling approach")

P("The chamber is modelled as a **lumped thermal network**: a small number of "
  "bodies, each at one temperature, connected by conductances, each storing "
  "energy in a capacitance. This is the standard resistor-capacitor network of "
  "circuit theory with the labels changed, and the analogy is exact rather than "
  "decorative:")

TBL(["Thermal", "Electrical", "Unit"],
    [["Temperature T", "Voltage", "K"],
     ["Heat flow Q̇, power P", "Current", "W"],
     ["Thermal conductance G, UA", "Conductance", "W/K"],
     ["Thermal resistance R_{th} = 1/G", "Resistance", "K/W"],
     ["Thermal capacitance C = m·c", "Capacitance", "J/K"],
     ["First law at a node", "Kirchhoff's current law", "—"],
     ["Time constant τ = C/G", "τ = RC", "s"]],
    weights=[2.4, 2, 1],
    caption="Table 9. The thermal-electrical analogy used throughout. The "
            "consequence worth stating: the whole model of §5.3 can be simulated "
            "unchanged in SPICE, with watts as current sources and kelvin above "
            "ambient as node voltages.")

P("Three composition rules follow from the first law and are used everywhere. "
  "Paths that share the same temperature difference are in parallel and their "
  "conductances add; stages traversed in sequence are in series and their "
  "resistances add; a coolant stream entering at one temperature and leaving at "
  "the node's temperature behaves as a conductance to the inlet, "
  "G_{adv} = ṁ·c_p.")

P("All temperatures are written as **deviations from ambient**, θ = T − T_a, so "
  "ambient sits at zero and the ambient-drift disturbance appears explicitly as "
  "a −Ṫ_a term in every row rather than being scattered through the constants.")

H(2, "Is lumping legitimate? — the Biot audit")

P("A body may be represented by a single temperature when its internal "
  "conduction resistance is small compared with its surface film resistance, "
  "i.e. when the Biot number is small:")

EQ("Bi = h·L_c / k = R_{internal} / R_{film} ,   lump valid for Bi < 0.1", tag="4")

P("Every element inside the boundary was checked against this criterion, with "
  "deliberately high film coefficients so the test is conservative "
  "(fan-stirred air h = 50 W/m²K, water side 3 000, TIM contact 5 000).")

TBL(["Element", "Path checked", "Bi", "Verdict"],
    [["Trim-heater copper foils", "35 µm Cu", "4 × 10⁻⁶", "Lump ✓"],
     ["Barrel resistor, ceramic core", "Ø20 mm cylinder", "0.010", "Lump ✓"],
     ["Pickup-radiator fins", "0.1 mm fin", "2 × 10⁻⁵", "Lump ✓"],
     ["Radiator tube wall", "0.5 mm brass, water side", "0.014", "Lump ✓"],
     ["Wiring, Cu core", "Ø1 mm", "3 × 10⁻⁵", "Lump ✓"],
     ["Fans, plastic impeller", "~2 mm PBT", "0.20",
      "Borderline — C ≈ 30–60 J/K, ~1 % of the total, and metered anyway"],
     ["Sensor PCBs", "1.6 mm FR-4", "0.13",
      "Borderline — C ≈ 2–5 J/K, negligible"],
     ["Chamber and gap air", "Fluid — Biot does not apply", "—",
      "Lumped **by stirring**: the fans' turnover time is seconds against "
      "minutes of system τ. This, not heat transfer, is the fans' real job — it "
      "is what makes T_i and T_e mean something"],
     ["RTD / TMP117 probes", "Sheath, package", "0.15–0.17",
      "Not lumps — modelled as first-order sensor lags 1/(τ_s·s + 1)"],
     ["**XPS shells, 50 mm**", "Foam slab", "**≈ 70**",
      "**Wall, not lump — by design.** Bi ≫ 1 is the definition of a wall: the "
      "temperature drop lives inside the material, so it is modelled as a "
      "conductance (G_{gap}, G_{out})"],
     ["**DUT board**", "1.6 mm FR-4 into air or TIM", "**0.27 … 27**",
      "**Fails — and that is why the instrument is a calorimeter.** Junctions run "
      "tens of kelvin above the baseplate, so the DUT can never be one "
      "temperature. The reading never uses a DUT temperature: at dT/dt = 0 the "
      "internal gradients are constant and every dissipated watt must still leave "
      "along a metered path, however ugly the inside is"]],
    weights=[1.8, 1.7, 0.9, 4.4], size=8.5,
    caption="Table 10. Biot audit. The two failures are structural features, not "
            "modelling errors: the foam is supposed to be a wall, and the DUT is "
            "supposed to be opaque to the measurement.")

H(2, "The node equations")

P("Three storage nodes survive the audit, plus one algebraic readout. With "
  "θ ≡ T − T_a throughout:")

EQ("C_i·dθ_i/dt = (1−α)·P_{DUT} + P_{aux} − UA_{rad}·(θ_i − θ_w) "
   "− G_{gap}·(θ_i − θ_e) − C_i·Ṫ_a", tag="5a")
EQ("C_w·dθ_w/dt = α·P_{DUT} + UA_{rad}·(θ_i − θ_w) − ṁc_p·(θ_w − θ_{in}) "
   "− C_w·Ṫ_a", tag="5b")
EQ("C_e·dθ_e/dt = P_e + G_{gap}·(θ_i − θ_e) − G_{out}·θ_e − C_e·Ṫ_a", tag="5c")
EQ("P_{meas} = ṁc_p·(θ_w − θ_{in}) − P_{aux}", tag="5d")

TBL(["Symbol", "Meaning", "Unit"],
    [["θ_i", "Chamber air, above ambient", "K"],
     ["θ_w", "Water side — resident loop water, radiator core, DUT mass", "K"],
     ["θ_e", "Guard gap air", "K"],
     ["θ_{in}", "Water inlet temperature from the reject loop", "K"],
     ["P_{DUT}", "The measurand — the DUT's dissipation", "W"],
     ["P_{aux}", "Metered auxiliaries inside B1 — chamber fans, preheat resistor", "W"],
     ["P_e", "Guard heater — the null actuator", "W"],
     ["α", "Fraction of the DUT's loss reaching the water directly rather than "
      "via the chamber air. **α = 0 for the built rig**", "—"],
     ["C_i, C_w, C_e", "Capacitances of the three nodes", "J/K"],
     ["UA_{rad}", "Chamber air ↔ water bridge through the pickup radiator", "W/K"],
     ["G_{gap}", "Chamber → gap. **The leak**", "W/K"],
     ["G_{out}", "Gap → ambient. Unmetered, and harmless", "W/K"],
     ["ṁc_p", "Advection through the water loop — the meter and the drain", "W/K"],
     ["Ṫ_a", "Ambient drift rate", "K/s"]],
    weights=[1.2, 5, 0.8],
    caption="Table 11. Symbols in the node equations.")

NOTE("Three derivatives, three states",
     "Only equations (5a)–(5c) carry a d/dt, so the model has exactly three "
     "states. Equation (5d) has none — it is a pure algebraic readout and belongs "
     "in the output matrix, never in the dynamics. This distinction is the one "
     "most often collapsed by mistake, and getting it wrong makes the reading "
     "appear to have dynamics it does not have.")

P("One identity in (5b) and (5d) is worth pointing at, because it is why the "
  "instrument works at all: **the water node's drain term and the meter's output "
  "are the same signal.** ṁc_p·(θ_w − θ_{in}) is simultaneously the plant's "
  "advection loss and the reported reading. That is also why θ_w must be a "
  "state — the reading is a capacitor voltage being observed, not an unknown "
  "being solved for.")

H(2, "The conductances")

P("Each conductance is a series/parallel assembly of the standard mechanisms.")

EQ("G_{gap} = [ d_{in}/(λ·A_{wall}) + 1/((h_{gap} + h_r)·A_{wall}) ]⁻¹ "
   "+ G_{seam} + G_{ft} + G_{wire}", tag="6")
EQ("G_{out} = [ d_{out}/(λ·A_{out}) + 1/((h_{ext} + h_{r,ext})·A_{out}) ]⁻¹", tag="7")
EQ("UA_{rad} = [ 1/(η_o·h_{air}·A_{air}) + R_{wall} + 1/(h_w·A_{water}) ]⁻¹", tag="8")
EQ("h_r = 4·ε_{eff}·σ·T̄³ ,    ε_{eff} = (1/ε_1 + 1/ε_2 − 1)⁻¹", tag="9")

P("Equation (9) is the linearised radiation term, obtained by factoring "
  "T_1⁴ − T_2⁴ = (T_1² + T_2²)(T_1 + T_2)(T_1 − T_2) and freezing the bracket at "
  "4T̄³. Over a 20–30 K band the error is a few percent of an already small term, "
  "and it is what keeps the whole model linear.")

NOTE("The foil rule — the single most load-bearing material choice",
     "Radiation across the gap runs skin-to-skin, in parallel with the foam's "
     "conduction, and it is the larger of the two if the skins are left bare. At "
     "T̄ = 320 K over the ≈0.6 m² gap: matte skins (ε ≈ 0.9 both sides) give "
     "ε_{eff} = 0.82, h_r ≈ 6.1 W/m²K and **G_{rad} ≈ 3.7 W/K** — 6.5× the "
     "foam's own 0.56 W/K. Aluminium foil on both facing skins (ε ≈ 0.05) gives "
     "ε_{eff} = 0.026 and **G_{rad} ≈ 0.11 W/K**, about 30× less. This matters "
     "twice over. A perfect null cancels radiation along with everything else, "
     "but the null is never perfect, and the leak error is G_{gap}·δT with "
     "G_{gap} *including* G_{rad} — so matte skins make every millikelvin of "
     "control error about seven times more expensive. Worse, the servo holds the "
     "gap **air** at T_i while radiation runs **skin to skin**, and the outer "
     "skin sits about 1 K below the gap air it feeds; with matte skins that alone "
     "leaks 3–4 W with a perfect air-side null, several times the entire budget. "
     "With foil, the same imperfection costs ~0.1 W.")

TBL(["Quantity", "Design value", "Where it comes from"],
    [["G_{gap}", "**0.56 W/K** ⇒ R_{th(in)} = 1.79 K/W",
      "XPS 50 mm, λ ≈ 0.035–0.04 W/m·K, A ≈ 0.6 m², foil-lined gap"],
     ["G_{out}", "**0.45 W/K**", "Outer shell, same class of construction"],
     ["UA_{rad}", "**30 W/K**",
      "PC-watercooling radiator with fans, air-side limited"],
     ["ṁc_p", "**3–30 W/K** design sweep (43–433 mL/min)",
      "Set per operating point from ṁ = P/(c_p·ΔT_{set})"],
     ["h_{gap}, h_{ext}", "5–15 W/m²K",
      "Fan-stirred and free air respectively — the softest numbers in the model, "
      "which is why commissioning measures the assembled G instead"],
     ["h_r, foil-lined", "≈ 0.16 W/m²K", "Equation (9) with ε_{eff} = 0.026"]],
    weights=[1.3, 2.2, 4],
    caption="Table 12. Conductance design values. Commissioning values override "
            "all of these.")

P("The ratio that makes the instrument work is **UA_{rad}/G_{gap} ≈ 54**. The "
  "metered path is about fifty times stiffer than the leak path, so the thermal "
  "current divider already sends the overwhelming majority of the heat into the "
  "water before the guard does anything; the guard then removes what the divider "
  "would still have lost. It is also why θ_i is pinned close to θ_w: a node tied "
  "to a rail through 30 W/K with only 0.56 W/K pulling elsewhere follows the rail.")

H(2, "The capacitances")

TBL(["Node", "Contents", "Design value", "Expected range"],
    [["C_i", "Chamber air (rigid box ⇒ c_v, not c_p), participating foam skin, "
      "fixtures, looms", "650 J/K", "0.3–1.0 kJ/K"],
     ["C_w", "Resident loop water inside B1, radiator core, DUT mass "
      "(coupled through the baseplate)", "1 800 J/K", "1.0–2.7 kJ/K"],
     ["C_e", "Gap air, both facing foam skins, foil, guard heater", "300 J/K",
      "0.1–0.5 kJ/K"],
     ["C_i + C_w", "**The mass the meter has to sweep out**", "**2 450 J/K**",
      "Sets the settling schedule and the steady gate"]],
    weights=[1, 3.6, 1.4, 1.8],
    caption="Table 13. Capacitance budget. Only the wall's inner slice "
            "participates on the run timescale — the participating depth is "
            "δ ≈ 2·√(a·t), saturating at the full thickness once the run is long "
            "enough.")

P("The foam wall's own distributed capacitance is real but is not one of the "
  "three states. It appears as an additional slow mode on top of the model, "
  "somewhere between 5.6 and 22 minutes depending on whether the slab is charged "
  "from one face or both — a factor of four that is still open and is the reason "
  "the first measurement point of a session takes hours while later points take "
  "minutes.")

H(2, "The error budget")

P("Relaxing both gates and keeping the deviations as first-order terms gives the "
  "error ledger. Each term is owned by a different part of the design:")

EQ("δP = G_{gap}·δT_{null} + C·|dT_i/dt| + c_{p,w}ΔT_w·δṁ + ṁc_{p,w}·δ(ΔT_w) "
   "+ δP_{aux} + δP_{par}", tag="10")

TBL(["Term", "At 15 W, ΔT_w = 5 K", "At 150 W, ΔT_w = 10 K", "Owner"],
    [["Flow calibration, 0.5 % gravimetric", "0.08 W", "0.75 W", "Calibration"],
     ["ΔT pair, 20 mK matched after co-calibration", "0.06 W", "0.30 W", "Sensing"],
     ["c_{p,w} property, 0.1 %", "0.02 W", "0.15 W", "Property data"],
     ["Auxiliary V·I metering, 0.5 %", "0.03 W", "0.05 W", "Metering"],
     ["Guard null residual, 0.3 K × G_{gap} (foil-lined)", "0.17 W", "0.17 W",
      "**Control loop**"],
     ["Storage, 0.2 W slice", "0.20 W", "0.20 W", "Steady-state detector"],
     ["Parasitics — anchored cabling, sealed seams", "0.05 W", "0.10 W",
      "Build quality"],
     ["**RSS total against P_{acc} = ±1 W**", "**≈ 0.29 W ✓**",
      "**≈ 0.87 W ✓ (tight)**", ""]],
    weights=[3.2, 1.5, 1.5, 1.5],
    caption="Table 14. Error budget at the two worst corners. The corners fail "
            "for different reasons: the cold corner is a null-plus-storage "
            "problem, owned by the control loops; the hot corner is a "
            "flow-calibration problem, owned by gravimetric rigour. Matte gap "
            "skins would put the null row at 1–4 W and sink both.")

P("Two properties of this budget are worth stating plainly. First, the "
  "**common-mode error of the ΔT pair cancels** — only the pair mismatch after "
  "co-calibration survives, so absolute RTD accuracy is far less important than "
  "matching. Second, **no control loop can fix a flow-calibration scale error**: "
  "if the firmware computes with a nominal ṁc_p while the loop drains at the true "
  "one, the reading is biased in proportion, and a 0.5 % flow error is 0.75 W at "
  "150 W regardless of how well the guard is tuned.")

# ═══════════════════════════════════════════ 6. state space & control
H(1, "State-Space Model and Guard Control Design")

P("The node equations of §5.3 are assembled into one state-space object. This "
  "is not a change of physics; it is the form that makes three questions "
  "answerable by inspection rather than by simulation: is the meter unbiased, "
  "how fast are the modes, and what does each error source cost in watts.")

H(2, "States, inputs and output")

EQ("x = [ θ_i , θ_w , θ_e ]ᵀ ,   u = [ P_{aux} , P_e ]ᵀ ,   "
   "d = [ P_{DUT} , θ_{in} , Ṫ_a ]ᵀ ,   y = P_{meas}")

P("The split between u and d is the control-theoretic one: u is what the "
  "firmware can command, d is what it must suffer. **P_{DUT} is a disturbance** "
  "— in control terms, a measurand that cannot be set is a disturbance, however "
  "central it is to the purpose of the instrument.")

EQ("ẋ = A·x + B_u·u + B_d·d ,    y = C·x + D_u·u + D_d·d", tag="11")

H(2, "The matrices")

P("Each row is one node's balance law divided by that node's own capacitance; "
  "each column is the state or input it is differentiated against. That is the "
  "whole derivation.")

TBL(["A", "θ_i", "θ_w", "θ_e"],
    [["**air**", "−(UA_{rad} + G_{gap})/C_i", "UA_{rad}/C_i", "G_{gap}/C_i"],
     ["**water**", "UA_{rad}/C_w", "−(UA_{rad} + ṁc_p)/C_w", "0"],
     ["**guard**", "G_{gap}/C_e", "0", "−(G_{gap} + G_{out})/C_e"]],
    weights=[1, 2, 2, 2],
    caption="Table 15. The system matrix A. Every entry has units of s⁻¹.")

TBL(["", "B_u : P_{aux}", "B_u : P_e", "B_d : P_{DUT}", "B_d : θ_{in}",
     "B_d : Ṫ_a"],
    [["**air**", "1/C_i", "0", "(1−α)/C_i", "0", "−1"],
     ["**water**", "0", "0", "α/C_w", "ṁc_p/C_w", "−1"],
     ["**guard**", "0", "1/C_e", "0", "0", "−1"]],
    weights=[1, 1.2, 1, 1.3, 1.3, 1],
    caption="Table 16. The input matrices. The Ṫ_a column is −1 in every row: "
            "that is the deviation coordinate system moving under the model, not "
            "heat flowing through a conductance.")

EQ("C = [ 0 , ṁc_p , 0 ] ,    D_u = [ −1 , 0 ] ,    D_d = [ 0 , −ṁc_p , 0 ]", tag="12")

P("The two non-zero D entries are the auxiliary-power subtraction and the "
  "inlet-temperature subtraction. They bypass the integrators entirely, which is "
  "why noise on P_{aux} or on the inlet RTD appears in the reading instantly and "
  "unfiltered — a direct consequence of one matrix entry, and a reason to meter "
  "the auxiliaries well.")

P("Note that ṁc_p appears in **four** places — inside A, inside B_d, inside C and "
  "inside D_d. It is one conductance with two ends: the water node's drain, the "
  "rail it drains to, and the meter that watches it.")

H(2, "Structural audit")

P("Seven checks are run on A before any number is trusted. They catch wiring "
  "errors — signs, missing bridges, bad couplings — that no amount of simulation "
  "makes obvious. All seven are coded as assertions in "
  "`calorimetry-ss.py: audit()` and run on every plant build.")

TBL(["#", "Check", "What it means", "Catches"],
    [["1", "All diagonal entries negative", "Every node decays on its own",
      "A sign error that makes the simulation diverge"],
     ["2", "Row 1 sums to zero",
      "The air has no exit except through the water and the guard", "A missing path"],
     ["3", "Row 2 sums to −ṁc_p/C_w",
      "The water leaks to the inlet rail — **this is the metered exit**",
      "A meter wired to the wrong rail"],
     ["4", "Row 3 sums to −G_{out}/C_e",
      "The guard leaks to ambient — **this is the unmetered exit, and the whole "
      "error term**", "A guard accidentally modelled as sealed"],
     ["5", "C_i·a_{12} = C_w·a_{21} = UA_{rad}; C_i·a_{13} = C_e·a_{31} = G_{gap}",
      "Each shared bridge appears twice, with opposite signs, weighted by its own "
      "capacitance", "Energy created at a coupling — the most common multi-node "
      "wiring error"],
     ["6", "a_{23} = a_{32} = 0", "Water and guard do not touch directly",
      "A spurious coupling"],
     ["7", "All eigenvalues real and negative",
      "A purely dissipative single-domain network cannot oscillate",
      "A sign error anywhere. **Only valid open-loop** — see §6.7"]],
    weights=[0.4, 3, 3.4, 3], size=8.5,
    caption="Table 17. The structural audit.")

NOTE("Two non-zero row sums means two ways out of the boundary",
     "One of them is measured; the other is the error term. That single sentence "
     "is the entire physics of a calorimeter, and it is legible in the row sums "
     "of A before a single number is entered.")

H(2, "Modes")

P("At the 150 W / 10 K operating point (ṁc_p = 15 W/K, 216 mL/min) the system "
  "matrix and its modes are:")

MONOBLOCK([
    "A =  [ -0.047015   0.046154   0.000862 ]",
    "     [  0.016667  -0.025000   0.000000 ]   s^-1",
    "     [  0.001867   0.000000  -0.003367 ]",
], size=9)

TBL(["Mode", "τ", "Eigenvector [θ_i, θ_w, θ_e]", "What it physically is"],
    [["Fast", "**15.2 s**", "[+1.000, −0.408, −0.030]",
      "Air and water **exchanging** across the radiator — opposite signs, guard "
      "uninvolved. In circuit terms, the two capacitors in series across the "
      "radiator bridge"],
     ["Mid", "**157.8 s**", "[+1.000, +0.893, −0.628]",
      "Air and water **moving together**, draining through the water bus. The two "
      "capacitors in parallel, discharging through the advection conductance"],
     ["Slow", "**314.5 s**", "[+0.100, +0.077, +1.000]",
      "The guard shell charging against ambient — almost entirely θ_e"]],
    weights=[0.9, 1, 2, 4.5], size=9,
    caption="Table 18. Modes at ṁc_p = 15 W/K. All three eigenvalues are real and "
            "negative, as they must be for a passive single-domain network.")

P("The eigenvectors are the point of the exercise. Three time constants alone "
  "say \"one fast thing, one medium thing, one slow thing\"; the eigenvectors say "
  "**which physical process each one is**, and immediately that the five-minute "
  "mode is the guard shell rather than the measurement. Note also that the modal "
  "time constants are not the per-node C/ΣG figures (21 s, 40 s, 297 s): the fast "
  "mode is faster than either node alone and the mid mode is slower than both. "
  "Coupled nodes have modes, not individual time constants.")

H(2, "The measurement, as a DC gain")

P("The whole instrument reduces to one question: **is the zero-frequency gain of "
  "the P_{DUT} → P_{meas} channel equal to one?** With the guard open, it is not.")

EQ("G(0) = −C·A⁻¹·B + D")

TBL(["ṁc_p [W/K]", "Flow [mL/min]", "DC gain, guard **open**",
     "Deficit ε", "Leakage-divider formula", "Agreement"],
    [["3.00", "43", "0.916183", "8.38 %", "0.083817", "1.6 × 10⁻¹⁵"],
     ["5.00", "72", "0.944985", "5.50 %", "0.055015", "9.2 × 10⁻¹⁶"],
     ["7.50", "108", "0.960076", "3.99 %", "0.039924", "8.2 × 10⁻¹⁶"],
     ["10.00", "144", "0.967804", "3.22 %", "0.032196", "3.8 × 10⁻¹⁶"],
     ["15.00", "216", "0.975657", "2.43 %", "0.024343", "1.0 × 10⁻¹⁶"],
     ["20.00", "289", "0.979631", "2.04 %", "0.020369", "2.6 × 10⁻¹⁶"],
     ["25.00", "361", "0.982032", "1.80 %", "0.017968", "4.3 × 10⁻¹⁶"],
     ["30.00", "433", "0.983638", "1.64 %", "0.016362", "1.0 × 10⁻¹⁶"]],
    weights=[1.2, 1.2, 1.6, 1, 1.6, 1.2],
    caption="Table 19. Systematic error with the guard open, α = 0. Computed two "
            "independent ways — from the state-space DC gain and from the closed "
            "form of equation (13) — and they agree to machine precision.")

P("The deficit has a closed form. Writing everything as resistances R = 1/G:")

EQ("ε = 1 − G(0) = [ R_{adv} + (1−α)·R_{rad} ] / "
   "[ R_{adv} + R_{rad} + R_{gap} + R_{out} ]", tag="13")

P("It is a **leakage divider**: the reading is short by the ratio of the metered "
  "path's resistance to the total resistance of all paths out of the boundary. "
  "Four consequences follow immediately:")

TBL(["Limit", "DC gain", "Why"],
    [["G_{gap} → 0 — a perfect wall, **or a guard that nulls θ_i − θ_e**",
      "**Exactly 1**", "No alternative exit exists"],
     ["G_{out} → 0 — the guard sealed from ambient", "**Exactly 1**",
      "Heat may enter the gap but cannot leave the boundary, so it must come back "
      "out through the water"],
     ["Any α, with either of the above", "**Exactly 1**",
      "**The meter is physically incapable of telling how the heat reached the "
      "water**"],
     ["Small G_{gap}", "1 − G_{gap}·(R_{adv} + (1−α)·R_{rad})",
      "The error is first order in the gap conductance — halve the leak, halve "
      "the error"]],
    weights=[3, 1.3, 4],
    caption="Table 20. Limits of the leakage divider.")

P("The third row is the α-invariance result, and it is what made the copper "
  "block optional rather than necessary: **the reading does not depend on how the "
  "heat got into the water.** Verified numerically — at 150 W / 10 K the "
  "open-guard gain differs between the two configurations (0.9757 at α = 0, "
  "0.9835 at α = 0.97) but the **closed-loop gain is exactly 1.00000000 in "
  "both**, with a residual error of order 10⁻¹² W, which is integrator "
  "convergence, not bias.")

H(2, "Guard nulled — the exact meter response")

P("Under the null gate (θ_e ≡ θ_i) the G_{gap} terms cancel on both sides and "
  "the plant collapses to a 2 × 2. Three DC gains then define the instrument:")

TBL(["Channel", "Value", "Meaning"],
    [["P_{DUT} → P_{meas}", "**1.000000**",
      "Exact, for every α, C_i, C_w, UA_{rad} and ṁc_p"],
     ["θ_{in} → P_{meas}", "**0**",
      "The reading is immune to the water setpoint — so auto-ranging the water ΔT "
      "between points is free"],
     ["P_{aux} → P_{meas}", "≈ 0 (machine precision)",
      "The D_u = −1 subtraction provably works. Delete that one entry and every "
      "watt of preheat is reported as DUT loss"]],
    weights=[1.8, 1.3, 4.5],
    caption="Table 21. The three gains that define the instrument, with the guard "
            "nulled.")

P("The transfer function of the meter, with the guard nulled, is second order "
  "with one zero:")

EQ("P_{meas}(s)/P_{DUT}(s) = [ 1 + α·(C_i/UA_{rad})·s ] / "
   "[ 1 + ((C_i+C_w)/ṁc_p + C_i/UA_{rad})·s + (C_i·C_w/(UA_{rad}·ṁc_p))·s² ]",
   tag="14", size=11)

P("Its two poles separate cleanly because UA_{rad} ≫ ṁc_p:")

EQ("τ_{fast} ≈ (1/UA_{rad})·(1/C_i + 1/C_w)⁻¹ ≈ 16 s ,    "
   "τ_{slow} ≈ (C_i + C_w)/ṁc_p")

P("τ_{slow} is the number that sets the schedule: it is the whole thermal mass "
  "inside the boundary being swept out through the meter. At 15 W and 43 mL/min "
  "that is 2 450/3 = 817 s, which is why the cold corner takes an hour to settle "
  "while the hot corner takes ten minutes. The zero sits at α·C_i/UA_{rad} — "
  "21 s in the block configuration, and it **disappears entirely at α = 0**. "
  "The split changes the transient; it never changes the answer.")

H(2, "The guard loop")

P("Collecting the θ_e terms of (5c) gives the SISO plant the guard controller "
  "actually sees:")

EQ("θ_e(s) = [ K_e/(τ_e·s + 1) ]·P_e(s) + [ β/(τ_e·s + 1) ]·θ_i(s)", tag="15")
EQ("K_e = 1/(G_{gap} + G_{out}) = 0.990 K/W ,   "
   "τ_e = C_e/(G_{gap} + G_{out}) = 297 s ,   "
   "β = G_{gap}/(G_{gap} + G_{out}) = 0.554")

NOTE("β = 0.554 is why the guard exists — and why proportional control is not enough",
     "Left unpowered, the gap settles at only 55 % of the chamber's temperature "
     "rise, because it loses the rest to the room through G_{out}. A pure "
     "proportional controller leaves a steady-state error of (1−β) divided by the "
     "loop gain, and that error is a permanent leak straight into the reading. "
     "The integrator is not a refinement — it is the accuracy. §7.8 prices it.")

P("The controller is tuned by the λ-rule (internal model control), which sets "
  "one knob — the desired closed-loop time constant λ_g — and derives both gains "
  "from the identified plant:")

EQ("K_{p,g} = τ_e/(K_e·λ_g) = 297/(0.990 × 60) = **5.000 W/K** ,   "
   "K_{i,g} = K_{p,g}/τ_e = **0.01683 W/(K·s)**", tag="16")

P("λ_g = 60 s is chosen as roughly a fifth of the open-loop τ_e — fast enough "
  "that the guard follows a moving chamber temperature, slow enough not to "
  "amplify sensor noise on a difference of two temperatures. Two nonlinearities "
  "sit on top and are mandatory:")

BUL([
    "**Saturation, [0, P_max].** The floor is **zero, not −P_max**: heaters "
    "cannot cool. After a load dump the guard has no authority at all and the gap "
    "must bleed off through G_{out} on its own, which is slow. This is a real "
    "operational limit, visible in Figures 5 and 6.",
    "**Anti-windup, by back-calculation.** The integrator is what makes the "
    "instrument exact, so the moment the heater clips, the thing that guarantees "
    "accuracy stops accumulating correctly unless it is explicitly unwound. The "
    "implementation adds (P_e − v)/T_t to the integrator rate, with "
    "T_t = K_p/K_i = τ_e, which is exactly zero whenever the output is not "
    "clipped.",
])

NOTE("The closed loop may oscillate — and that is not a fault",
     "The structural audit of Table 17 asserts that all eigenvalues are real, on "
     "the grounds that a purely dissipative single-domain network has nothing to "
     "trade energy with. That is true of the plant, and only of the plant. Once "
     "the guard PI is closed, the integrator is a second storage type — energy "
     "moves between the gap's thermal capacitance C_e and the controller's "
     "accumulated error, exactly as it does between L and C in a resonant "
     "circuit — and the closed loop legitimately acquires a complex pole pair. "
     "The audit must therefore never be run on the closed-loop matrix: complex "
     "eigenvalues in the **plant** mean a sign error; complex eigenvalues in the "
     "**closed loop** mean you have a controller. A damping ratio near 0.7 is "
     "well tuned; much below 0.4 says to slow λ_g down.")

P("An optional feedforward term, P_{ff} = G_{out}·(θ_{target}), can be added. It "
  "is free insurance: its DC gain to the reading's error is **exactly zero**, so "
  "a wrong G_{out}, a wrong ambient or a stale load estimate degrade the "
  "transient null error and nothing else. The integrator owns the steady null.")

H(2, "What is deliberately not modelled")

P("Two rules keep the model from growing without improving:")

BUL([
    "**Anything whose datasheet or measured rating is the decision property.** "
    "The radiator is a UA figure, not a fin geometry; the pump is a flow band and "
    "a steps-per-millilitre constant, not an impeller; the foam is a λ and a "
    "thickness, not a microstructure; radiation is a grey ε, not a spectrum. The "
    "justification is that commissioning measures the **assembled** C and G, so "
    "every sub-component detail is absorbed into two calibrated numbers — "
    "modelling it twice would add error, not fidelity.",
    "**Anything below the error floor.** The standing promotion rule: a neglected "
    "effect earns a state or an error term only when its worst case exceeds about "
    "10 % of P_{acc}, i.e. 0.1 W. Checked and dismissed on that basis: chamber-air "
    "c_p versus c_v, pump-head enthalpy, radiation through opaque foam, barometric "
    "drift, c_{p,w}(T) beyond evaluating at the mean, and the DUT's internal "
    "thermal structure.",
])

P("Four effects are genuinely nonlinear and are handled in the simulation rather "
  "than in the matrices: heater saturation, anti-windup, the bilinearity of ṁ "
  "(it multiplies a state, so the model is LTI only at a frozen flow — which is "
  "the operating protocol anyway), and sensor quantisation.")

# ═════════════════════════════════════════════════ 7. python simulation
H(1, "System Simulation in Python")

H(2, "Purpose and scope")

P("The simulation exists to answer questions the algebra cannot: what the "
  "settling schedule actually is, how large the guard heater has to be, how big "
  "the transient null error gets before the loop catches up, and what happens "
  "when the actuator saturates. Saturation and anti-windup are not linear, so "
  "they cannot live inside A — which is why the model is integrated numerically "
  "rather than solved with a step-response function. Inside the limits it is "
  "exactly the LTI model of §6; outside them it is a different one.")

H(2, "Files and how to run it")

TBL(["File", "Role"],
    [["`calorimetry-ss.py`",
      "The L1 plant. Builds A, B_u, B_d, C, D_u, D_d at a frozen flow; runs the "
      "structural audit of Table 17 as assertions; provides unit helpers "
      "(mL/min ↔ kg/s ↔ W/K) and the modal decomposition. Runnable on its own — "
      "it prints the modes, the open-guard DC gains and a corner sweep."],
     ["`calorimetry_sim.py`",
      "**The main script.** Wraps the guard PI, saturation and anti-windup around "
      "that plant, integrates the resulting 4-state system, rebuilds every probe "
      "signal from the trajectory, prints the probe map and writes the figures."],
     ["`params.yaml`",
      "Every parameter, with its kind tag. Nothing numeric is hard-coded in "
      "either script."]],
    weights=[1.4, 5.5],
    caption="Table 22. The three source files, in `~/roboteq/python/`.")

MONOBLOCK([
    "$ cd ~/roboteq/python",
    "$ python calorimetry_sim.py            # uses the defaults at the top of the file",
    "$ python calorimetry_sim.py 300 12     # or: P_DUT = 300 W, water dT = 12 K",
], size=9)

P("Dependencies are `numpy`, `matplotlib`, `pyyaml` and `python-control`. The "
  "script checks for all four against the interpreter that is actually running "
  "it and prints the exact `pip` line to fix a missing one, because the common "
  "failure here is having the packages installed against a different Python than "
  "the one being invoked. It also verifies that its output directory is writable "
  "before doing any work, and reports the size and wall-clock timestamp of every "
  "figure it writes — so that a figure that did not change can be distinguished "
  "from a viewer that is caching.")

H(2, "The parameter file")

TBL(["Parameter", "Value", "Kind", "Note"],
    [["`C_i`, `C_w`, `C_e`, `C_r`", "650, 1800, 300, 2000 J/K", "estimate",
      "Node capacitances; C_r is the reject-side reservoir"],
     ["`UA_rad`", "30 W/K", "estimate", "Air ↔ water bridge"],
     ["`G_gap`", "0.56 W/K", "estimate", "**The leak**"],
     ["`G_out`", "0.45 W/K", "estimate", "Gap → ambient"],
     ["`alpha`", "0.97", "design pick",
      "**Stale — the shipping configuration is 0.0.** See the note in §4.7"],
     ["`c_p_w`, `rho_w`", "4180 J/kg·K, 995 kg/m³", "property", "Water at ~305 K"],
     ["`mdot`", "0.00358852 kg/s", "estimate",
      "Nominal only. Not a free parameter — the supervisor derives it per range "
      "as ṁ = P/(c_p·ΔT_{set}). This value is the 150 W / 10 K corner"],
     ["`ranges`", "50 W → 5 K; 200 W → 10 K; 600 W → 25 K", "control",
      "The auto-ranging table"],
     ["`lambda_guard`", "60 s", "control", "Guard closed-loop time constant"],
     ["`P_e_max`", "100 W", "design pick", "Guard heater ceiling"],
     ["`P_acc`", "1.0 W", "control", "Accuracy target"],
     ["`P_thr`", "0.1 W", "control", "Residual threshold for the steady gate"],
     ["`sigma_pair_meter`, `sigma_pair_guard`", "20 mK each", "measurement",
      "Sensor-pair residuals after co-calibration"],
     ["`flow_cal_err`", "0.5 %", "calibration", "Gravimetric flow calibration"]],
    weights=[2.2, 2.2, 1.2, 4], size=8.5,
    caption="Table 23. `params.yaml`. Kind tags mark whether a number is a design "
            "choice, a material property, a control target, a sensor spec or a "
            "value that commissioning will replace.")

H(2, "Code structure")

TBL(["Function", "What it does"],
    [["`build_matrices(mcp)`",
      "Returns the six blocks at a frozen advection conductance. ṁ is bilinear — "
      "it multiplies a state — so the model is LTI only at a fixed flow, and a "
      "new plant is built per operating point."],
     ["`audit(A, mcp)`",
      "The seven structural assertions of Table 17. Runs on every plant build. "
      "Documented as **open-loop only**: the closed loop legitimately grows a "
      "complex pair, so running this on A_cl would fail for the wrong reason."],
     ["`build_plant(mcp)`",
      "Wraps the matrices as a named `python-control` state-space object, and "
      "asserts that ṁc_p is inside the design sweep — a units check that catches "
      "a flow entered in the wrong unit before it can produce plausible-looking "
      "nonsense."],
     ["`modes(A)`",
      "Eigenvalues as time constants, slowest first, each with its normalised "
      "eigenvector — Table 18."],
     ["`GuardPI`",
      "The controller: λ-tuned PI, output saturation with a floor of zero, and "
      "back-calculation anti-windup whose correction term is exactly zero when "
      "the output is not clipped. Gains are derived in `__init__` from the plant "
      "constants, not typed in."],
     ["`simulate(...)`",
      "Fixed-step RK4 over the 4-state system (3 plant states plus the guard "
      "integrator). Optionally pre-settles the rig with the DUT off for 30 000 s "
      "first, so t = 0 is a parked machine and the trace is a clean step response "
      "rather than a cold-start artefact. Returns a dictionary of named probe "
      "traces rebuilt from the stored trajectory."],
     ["`demo_windup(...)`",
      "Steps the load, then dumps it, at four actuator ceilings. The ceilings are "
      "**derived from the operating point** rather than hard-coded, so the "
      "demonstration keeps its meaning when the power or ΔT is changed."]],
    weights=[1.5, 5.5], size=9,
    caption="Table 24. The main entry points.")

H(2, "The probe map")

P("Every signal the model exposes is rebuilt after integration and reported as a "
  "named probe, with its peak and its final value. The list is deliberately the "
  "same set of signals that the firmware will expose in telemetry, so that "
  "simulation traces and bench traces can be compared line for line.")

TBL(["Probe", "Expression", "What it tells you", "Unit"],
    [["`P_meas`", "ṁc_p·(θ_w − θ_{rail}) − P_{aux}", "The reported reading", "W"],
     ["`eps`", "P_{meas} − P_{DUT}",
      "**The error trace — the reason to build the model at all**", "W"],
     ["`e`", "θ_i − θ_e", "The null error the guard controller acts on", "K"],
     ["`P_leak`", "G_{gap}·e", "The leak, in watts", "W"],
     ["`P_e`", "Controller output after the limiter",
      "Actuator demand against its ceiling", "W"],
     ["`v`", "Controller output before the limiter",
      "Saturation shows up as v ≠ P_e", "W"],
     ["`P_rad`", "UA_{rad}·(θ_i − θ_w)", "Watts crossing the pickup radiator", "W"],
     ["`P_adv`", "ṁc_p·(θ_w − θ_{rail})", "Watts carried by the flow", "W"],
     ["`P_out`", "G_{out}·θ_e", "Watts lost to ambient — unmetered", "W"],
     ["`dT_w`", "θ_w − θ_{rail}", "The auto-range variable", "K"],
     ["`th_i`, `th_w`, `th_e`", "States", "Node temperatures above ambient", "K"],
     ["`x_I`", "Guard integrator state", "Windup detector", "K·s"],
     ["`sat_hi`, `sat_lo`", "Booleans", "Pinned at the ceiling / at zero", "—"]],
    weights=[1.4, 2.2, 4, 0.7], size=8.5,
    caption="Table 25. The probe map.")

H(2, "Simulation results")

P("The figures below are the direct output of `calorimetry_sim.py`. All were run "
  "in the block-engaged configuration (α = 0.97), which is the script's default; "
  "as established in §6.5 this changes the **transient shape** but leaves the "
  "steady-state reading exactly unchanged, so the conclusions transfer to the "
  "built α = 0 rig. §7.7 gives the shipping-configuration numbers.")

P("Each response figure has the same six panels:")

TBL(["Panel", "Shows", "What to look for"],
    [["**Top left** — the measurement",
      "P_{DUT} (dashed, the truth) against P_{meas} (solid, the reading)",
      "That the reading converges onto the truth, and how long that takes"],
     ["**Top centre** — eps = reading − truth",
      "The error trace, with the ±1 W target marked. The y-axis is clipped "
      "because the t = 0 transient runs far off scale",
      "The minute at which the error first enters the ±1 W band, and that it then "
      "goes to zero rather than to a small constant"],
     ["**Top right** — node temperatures",
      "θ_i (air), θ_w (water), θ_e (guard) and the inlet rail, all in kelvin "
      "above ambient",
      "That θ_e tracks θ_i — the guard doing its job — and where the chamber "
      "settles relative to the rail"],
     ["**Bottom left** — guard null error",
      "e = θ_i − θ_e against the 1.79 K null gate",
      "The transient overshoot, whether it breaches the gate, and how fast it "
      "returns"],
     ["**Bottom centre** — actuator",
      "Demand v (dashed) against delivered P_e (solid), with P_max marked",
      "Saturation, which shows as the two curves separating. Also the steady "
      "guard power, which is what sizes the heater"],
     ["**Bottom right** — where the watts go",
      "P_rad (air→water), P_adv (the reading), P_leak (air→gap), P_out "
      "(gap→ambient)",
      "That P_adv converges to the DUT power, that P_leak collapses to zero, and "
      "that P_out is non-zero but harmless"]],
    weights=[1.9, 3.2, 3.6], size=8.5,
    caption="Table 26. How to read the six-panel response figures.")

H(3, "Nominal operating point — 150 W at 10 K water rise")

FIG("response_150W_dT10K.png",
    "**Closed-loop step response at the nominal point: P_DUT = 150 W, "
    "ΔT_set = 10 K, ṁc_p = 15 W/K (216 mL/min).** The reading reaches 1 % of "
    "final in **12.9 min** and enters the ±1 W band at about 14 min. The guard "
    "null error peaks at **1.76 K**, just under the 1.79 K gate, and the leak "
    "peaks at 0.99 W before collapsing. Steady guard power is **6.8 W**, against "
    "a 100 W ceiling — the actuator never comes close to saturating, and demand "
    "and delivered power are indistinguishable in the bottom-centre panel. In the "
    "bottom-right panel, note that P_rad goes sharply **negative** in the first "
    "seconds: the water starts at the inlet rail temperature, below the chamber, "
    "so heat initially flows from water to air until the radiator bridge "
    "equalises. Note also that P_out settles to exactly the same value as the "
    "guard heater power — in steady state, with the leak nulled, the guard heater "
    "supplies precisely what the gap loses to the room, and none of it touches "
    "the reading.")

H(3, "Same power, lower flow — the flow/settling trade")

P("Running the same 150 W at a larger water ΔT means less flow, which is "
  "attractive for the ΔT-pair error budget but slow, because τ_slow = "
  "(C_i + C_w)/ṁc_p. Figures 2 and 3 are the same load at 15 K and 20 K.")

FIG("response_150W_dT15K.png",
    "**150 W at ΔT_set = 15 K, ṁc_p = 10 W/K (144 mL/min).** Settling stretches "
    "to **19.1 min** and the null error peaks at **1.98 K**, now above the "
    "1.79 K gate. Steady guard power rises to 9.1 W because the chamber sits "
    "higher above ambient (θ_i ≈ 20 K rather than 15 K), so the gap loses more "
    "through G_out and the guard has to replace it.")

FIG("response_150W_dT20K.png",
    "**150 W at ΔT_set = 20 K, ṁc_p = 7.5 W/K (108 mL/min).** Settling **25.4 "
    "min**, null-error peak **2.12 K**, steady guard power **11.3 W**. Comparing "
    "Figures 1–3 gives the trade in one line: **halving the flow doubles the "
    "settling time, raises the transient null error and raises the guard's steady "
    "power, while improving the ΔT-pair error by the same factor of two.** The "
    "auto-ranging rule in `params.yaml` exists to place this trade per operating "
    "point rather than fixing it once.")

H(3, "Top of the envelope — 200 W")

FIG("response_200W_dT8K.png",
    "**200 W at ΔT_set = 8 K, ṁc_p = 25 W/K (361 mL/min).** The fastest corner "
    "in the set: **8.0 min** to 1 %, inside the ±1 W band by about 9 min. Guard "
    "power is only **5.9 W** steady, because the high flow pins the chamber lower "
    "(θ_i ≈ 13 K above ambient). The null error still peaks at 1.95 K, above the "
    "gate — the transient overshoot is driven by the mismatch between the fast "
    "chamber and the slow guard shell, not by the power level, and it is present "
    "at every corner.")

NOTE("The transient null error breaches the gate at and above ~100 W",
     "This is a genuine finding of the simulation and does not appear anywhere in "
     "the hand analysis, because it only exists once the guard is a real loop "
     "with a real τ_e instead of an assumed perfect null. The mechanism is a "
     "20-to-1 speed mismatch: the chamber's fast mode is ~15 s while the guard "
     "shell's open-loop τ_e is 297 s, so the guard is chasing a reference that "
     "moves far faster than it can follow. **It is a transient, so it does not "
     "bias a logged point** — the null gate holds the reading until e is back "
     "inside 1.79 K. What it costs is schedule, and what it forbids is logging "
     "early. The candidate fix is to ramp the load and the feedforward together "
     "rather than stepping them, which is a sequencer change, not a model change.")

H(3, "Actuator sizing, saturation and anti-windup")

P("The last pair of figures answers a hardware question: how big does the guard "
  "heater have to be, and what happens if it is undersized. The load is stepped "
  "to the corner power, then **dumped to 15 W** partway through — the error "
  "reverses sign at that moment, which is precisely when a wound-up integrator "
  "shows itself. Four configurations are overlaid: the heater as designed "
  "(100 W), a marginal ceiling between the steady and peak demand with "
  "anti-windup on and off, and a deliberately undersized one below the steady "
  "demand.")

FIG("windup_150W_dT20K.png",
    "**Saturation and anti-windup at 150 W / 20 K, load dumped to 15 W at "
    "25 min.** Steady guard demand at this corner is 11.5 W with a transient peak "
    "of 16.6 W. **Blue, 100 W as designed:** never saturates, null error stays "
    "small, recovery after the dump takes 4.2 min. **Orange and green, 13.3 W "
    "marginal:** clipped 24 % of the run; the two traces are nearly "
    "indistinguishable, so with a ceiling this close to the demand, anti-windup "
    "is barely exercised (peak integrator 1 442 vs 1 473 K·s, recovery 5.2 vs "
    "5.3 min). **Red, 8.0 W undersized:** clipped 48 % of the run, the integrator "
    "winds to 6 418 K·s — more than eight times the healthy case — the null error "
    "swings ±6 K, and recovery takes **22.4 min** against 4.2. That is a sizing "
    "failure, and no amount of controller tuning fixes it.")

FIG("windup_200W_dT8K.png",
    "**The same study at 200 W / 8 K.** Steady demand 5.9 W, transient peak "
    "14.1 W — note that the steady demand is *lower* here than at 150 W / 20 K, "
    "because the high flow keeps the chamber cooler, while the transient peak is "
    "similar. An undersized 4.2 W heater is clipped for 58.7 % of the run and "
    "takes 17.4 min to recover. **The design conclusion is that the guard heater "
    "must be sized against the transient peak, not the steady demand** — roughly "
    "a factor of 2.5 apart at these corners. The specified 100 W barrel resistor "
    "gives about 6× margin over the worst transient peak seen, which also covers "
    "its second role as the guard-side preheat element.")

NOTE("Why every configuration has a slow tail after the load dump",
     "In all four traces the delivered power sits at exactly zero for several "
     "minutes after the dump, and the null error takes minutes to come back. That "
     "is not a controller problem: **heaters cannot cool.** Once the demand goes "
     "negative the output is pinned at the floor and the gap can only bleed its "
     "stored energy off through G_out, at its own τ. The consequence for the test "
     "sequence is that **stepping the load down is slower than stepping it up**, "
     "so a load map should be swept upward wherever possible.")

H(2, "Corner sweep — the shipping configuration")

P("The table below is the model run at **α = 0, the built rig**, across the "
  "design envelope. It is the single most useful output of the simulation, "
  "because it is what sets the test schedule.")

TBL(["P_{DUT}", "ΔT_{set}", "ṁc_p", "Flow", "Gain, guard open",
     "Gain, guard closed", "1 % settling", "Peak |e|", "Peak leak", "P_e steady"],
    [["15 W", "5 K", "3.00 W/K", "43 mL/min", "0.9162 (−8.4 %)", "**1.000000**",
      "**63.6 min**", "0.32 K ✓", "0.18 W", "4.7 W"],
     ["50 W", "8 K", "6.25 W/K", "90 mL/min", "0.9540 (−4.6 %)", "**1.000000**",
      "30.8 min", "1.07 K ✓", "0.60 W", "6.6 W"],
     ["100 W", "10 K", "10.0 W/K", "144 mL/min", "0.9678 (−3.2 %)", "**1.000000**",
      "19.4 min", "2.13 K ✗", "1.19 W", "8.3 W"],
     ["150 W", "10 K", "15.0 W/K", "216 mL/min", "0.9757 (−2.4 %)", "**1.000000**",
      "13.2 min", "3.16 K ✗", "1.77 W", "9.0 W"],
     ["200 W", "10 K", "20.0 W/K", "289 mL/min", "0.9796 (−2.0 %)", "**1.000000**",
      "10.1 min", "4.18 K ✗", "2.34 W", "9.8 W"]],
    weights=[1, 0.9, 1, 1.1, 1.5, 1.3, 1.2, 1, 1, 1], size=8.5,
    caption="Table 27. Closed-loop corner sweep at α = 0, flow matched to the "
            "load, step from a parked rig. ✓/✗ marks the 1.79 K null gate on the "
            "**transient** peak; every corner satisfies it in steady state, which "
            "is what the logged point depends on.")

P("Three readings, two of which are warnings:")

BUL([
    "**The closed-loop gain is exactly 1 at every corner.** Five different flows, "
    "five different plants, one answer. This is the central result of the whole "
    "modelling effort: the instrument is unbiased by construction, not by "
    "calibration.",
    "**The cold corner is an hour.** 63.6 min at 15 W. The mechanism is not the "
    "foam — it is that (C_i + C_w) ≈ 2.45 kJ/K has to be swept out through only "
    "3 W/K of advection at 43 mL/min. Three levers exist, cheapest first: raise "
    "the flow and accept a smaller water ΔT; reduce C_w (every kJ/K removed is "
    "worth roughly five minutes); or estimate the steady value from the transient "
    "instead of waiting for it, which is a firmware change rather than a hardware "
    "one and is discussed in §7.9.2.",
    "**The transient null error grows with power** and crosses the gate above "
    "about 100 W. As established above, it costs schedule rather than accuracy.",
])

H(2, "Ablation — what the guard is actually worth")

P("Same plant, same operating point, three guard configurations:")

TBL(["Guard configuration", "DC gain P_{DUT} → P̂", "Deficit", "Error at 150 W",
     "Steady null error e_{ss}", "G_{gap}·e_{ss}"],
    [["**PI — integral action**", "**1.000000**", "0.0000 %", "**0.000 W**",
      "0.00000 K", "0.0000 W"],
     ["P-only (K_i = 0)", "0.994433", "0.5567 %", "0.835 W", "1.49125 K",
      "0.8351 W"],
     ["**Guard OFF (open loop)**", "0.967543", "3.2457 %", "**4.869 W**",
      "8.69397 K", "4.8686 W"]],
    weights=[2, 1.5, 1, 1.2, 1.4, 1.2],
    caption="Table 28. Guard ablation at 150 W / 10 K, α = 0. The identity "
            "ε = G_{gap}·e_{ss} holds exactly in the last two rows, which is a "
            "strong consistency check: the error in watts really is the null "
            "error in kelvin multiplied by the leak conductance.")

P("Two conclusions, and the second is the less obvious one:")

NUM([
    "**The guard is load-bearing, not a refinement.** Without it the rig is a "
    "3.2 % instrument — 4.9 W at 150 W, five times outside the ±1 W target, and "
    "worse at low flow, where Table 19 shows the open-guard deficit reaching "
    "8.4 %. This is the same failure mode Nair (RD-03) measured on a real "
    "unguarded build.",
    "**What makes it exact is the integrator, not the hardware.** Proportional "
    "control alone recovers 83 % of the deficit, from 4.87 W down to 0.84 W — but "
    "0.84 W of the ±1 W budget is still spent on a single error term. The last "
    "0.84 W is bought by one state, the integrator, and by nothing else. \"Add "
    "integral action to the guard PI\" and \"make the calorimeter unbiased\" are "
    "the same sentence — which is also why anti-windup is not optional.",
])

H(2, "Findings and open items")

H(3, "What the model settles")

TBL(["#", "Result"],
    [["1", "The reading is **exactly unbiased** with the guard closed — DC gain "
      "1.000000 at every flow, every power and every α tested."],
     ["2", "The systematic error with the guard open is a **leakage divider**, "
      "equation (13), verified against the state-space DC gain to machine "
      "precision across the flow range."],
     ["3", "**The reading is independent of α.** The copper block was never an "
      "accuracy feature, which is what made dropping it a pure scope decision."],
     ["4", "**The guard's integral action is what makes the instrument exact** — "
      "3.25 % open, 0.56 % proportional-only, exactly zero with the integrator."],
     ["5", "The water setpoint and the metered auxiliary power both have "
      "**exactly zero** DC gain to the error, so auto-ranging the water ΔT "
      "between points is free and the P_aux subtraction is provably correct."],
     ["6", "**Guard heater sizing is set by the transient peak, not the steady "
      "demand** — about 2.5× apart. The specified 100 W part has ~6× margin."],
     ["7", "The settling schedule across the envelope: **10 min at 200 W to "
      "64 min at 15 W**, driven by (C_i + C_w)/ṁc_p."],
     ["8", "The steady gate can be **derived rather than guessed**: the slope "
      "threshold is the residual power threshold divided by the slowest observed "
      "time constant, computed per point."]],
    weights=[0.4, 9], size=9)

H(3, "Open items")

TBL(["#", "Item", "Why it matters", "Where it lands"],
    [["1", "**The revised 30–40 K target and the flow floor.** Treating the "
      "50–60 °C chamber target directly as the water-loop ΔT puts the 15 W corner "
      "at 5–7 mL/min. Simulated, that corner has an open-guard gain of only "
      "0.60–0.66, and 1 % settling of **6–8 hours**. It also sits below the "
      "model's own sanity window (`build_plant` asserts 1 < ṁc_p < 50 W/K, i.e. "
      "14–721 mL/min).",
      "The 30–40 K figure is a **chamber-above-ambient** target being used as a "
      "stand-in for the water-loop ΔT, and the two only converge if the pickup "
      "radiator is very effective. If it is applied literally to the water loop "
      "at low power, the cold corner becomes impractical.",
      "**Recommend keeping the auto-ranging table** (5 K at 15 W, rising to "
      "10 K by 200 W) and treating 50–60 °C as the chamber setpoint, not the "
      "water ΔT. Confirm against the first commissioning run."],
     ["2", "Transient null error breaches the 1.79 K gate above ~100 W",
      "Costs schedule and forbids early logging; new at closed loop",
      "Candidate fix is a ramped load and feedforward instead of a step — a "
      "sequencer change"],
     ["3", "The foam wall's boundary condition — charged from one face or both",
      "A factor of four (5.6 vs 22 min) on the first-point wait, which is on the "
      "critical path", "A short Python job, not yet done"],
     ["4", "`params.yaml` still defaults to `alpha: 0.97`",
      "Harmless to the answer, misleading to a reader", "One-line change"],
     ["5", "Model-based early stopping",
      "Estimating the steady value from the transient collapses the cold-corner "
      "wait by roughly an order of magnitude, for about 40 lines of firmware and "
      "no hardware",
      "Should be commissioned **alongside** the full wait for one campaign and "
      "compared, not shipped unvalidated"],
     ["6", "Monte-Carlo over the estimated parameter ranges",
      "Every number in this report is a point estimate from mid-range design "
      "values", "Turns each figure into a confidence interval"],
     ["7", "Sensor lags τ_s are placeholders (5 s)",
      "They set the effective dead time the PI is tuned against", "Commissioning"],
     ["8", "Zonal guard — the gap is modelled as one lump",
      "A zero *average* null is not a zero leak; patchy faces still pay "
      "ΣG_face·δT_face", "After single-zone tuning works"]],
    weights=[0.4, 3, 3.2, 3], size=8.5,
    caption="Table 29. Open items, in priority order.")

NOTE("Commissioning overrides everything in this report",
     "Every conductance and capacitance above is a design estimate computed from "
     "geometry and material data. The commissioning step test — inject a known "
     "power with the calibration resistor, take C from the initial slope "
     "(C = P/Ṫ|₀), G from the endpoint (G = P/(T_ss − T_a)), and check that τ = "
     "C/G is consistent — measures the **assembled** rig and replaces them. The "
     "equations, the matrix structure, the audit and every conclusion in "
     "sections 5 to 7 survive that substitution unchanged; only the numbers move.")

# ═══════════════════════════════════════════════════ 8 & 9 — placeholders
H(1, "Hardware Design")

NOTE("Section to be completed",
     "This section will cover the carrier PCB for the ESP32-S3 DevKitC: the "
     "12 V → 5 V buck converter, the two independent heater power domains (inner "
     "chamber and guard) with their MOSFET stages and current sensing, the "
     "stepper driver for the peristaltic pump, PWM connectors for the fans, the "
     "RTD signal conditioning, and the protection strategy — ESD, overcurrent, "
     "overvoltage, short-circuit and thermal. It will also carry the bottom-up "
     "rail budget (≈22 A on the 12 V rail against a 29 A supply) and the "
     "measurement-grade versus protection-grade split on the current-sense "
     "channels.")

H(1, "Firmware")

NOTE("Section to be completed",
     "This section will cover the Zephyr application on the ESP32-S3: the task "
     "split across the control, sensing, supervisor and telemetry threads; the "
     "devicetree description of the signal chains; the discretised guard loop "
     "(1 s sample rate, plant discretised zero-order-hold, controller "
     "Tustin) with its saturation and anti-windup; the null and steady gates of "
     "§4.6.3; the auto-ranging supervisor state machine; the staged power-up and "
     "heater interlocks; and the telemetry path — one JSON snapshot per second "
     "over WebSocket, rendered both by the local LVGL panel and by a browser "
     "dashboard.")

# ─────────────────────────────────────────────────────────────── save
OUT = SRC
BAK = SRC.with_name(SRC.stem + " — pre-report backup.docx")
if not BAK.exists():
    shutil.copy2(SRC, BAK)
    print(f"backup written: {BAK}")
doc.save(str(OUT))
print(f"saved: {SRC}")
print(f"figures inserted: {FIGN[0]}")
